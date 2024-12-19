import os
import re
import uuid
import json
import logging
from datetime import datetime
import requests
from flask import Flask, render_template, request, send_file, jsonify, make_response
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename

# Configure logging
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                   handlers=[logging.FileHandler('app.log'), logging.StreamHandler()])

app = Flask(__name__)
logger = app.logger

# Configuration
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'pdf'}

# Ensure required directories exist
for folder in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
    os.makedirs(folder, exist_ok=True)

def allowed_file(filename):
    """Check if the file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_filename(filename):
    """Generate a clean, secure filename"""
    base = secure_filename(os.path.splitext(filename)[0])
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{base}_{timestamp}"

def extract_questions_from_pdf(pdf_path):
    """Extract individual questions from PDF with improved pattern matching"""
    questions = []
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = ''
            
            # Extract text from each page
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:  # Check if text extraction was successful
                    full_text += page_text + '\n'
        
        # Enhanced pattern matching for questions
        patterns = [
            r'(?:^|\n)(?:Q|Question)[.\s]*\d+[.)]\s*([^\n]+(?:\n(?!\d+[.)]\s*)[^\n]+)*)',
            r'(?:^|\n)\d+[.)]\s*([^\n]+(?:\n(?!\d+[.)]\s*)[^\n]+)*)',
            r'(?:^|\n)(?:[A-Z])\.\s*([^\n]+(?:\n(?![A-Z]\.\s*)[^\n]+)*)'
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, full_text, re.MULTILINE)
            for match in matches:
                question = re.sub(r'\s+', ' ', match.group(1).strip())
                if len(question) > 10:  # Minimum length check
                    questions.append(question)
        
        # Remove duplicates while preserving order
        questions = list(dict.fromkeys(questions))
        
        logger.info(f"Successfully extracted {len(questions)} questions from PDF")
        return questions
    except Exception as e:
        logger.error(f"Error extracting questions from PDF: {str(e)}")
        return []

def extract_text_from_response(api_response):
    """Extract and clean text from Gemini API response"""
    try:
        candidates = api_response.get('candidates', [])
        if candidates:
            answer = candidates[0].get('content', "").strip()
            # Clean up the answer
            answer = re.sub(r'\n{3,}', '\n\n', answer)  # Remove excessive newlines
            answer = answer.replace('```', '')  # Remove code block markers
            return answer
        return "No answer could be generated for this question."
    except Exception as e:
        logger.error(f"Error extracting API response: {str(e)}")
        return "Error processing the answer."

def get_answer_from_gemini(question):
    """Get answer from Gemini API with enhanced prompting"""
    API_KEY = "AIzaSyDpKtMhIzD-ThtEt6dr4InKQvuQJkcsroc"  # Replace with your actual API key
    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={API_KEY}"
    
    prompt = f"""Provide a comprehensive academic answer to the following question. 
    Focus on accuracy, clarity, and completeness while maintaining a professional tone:
    
    {question}
    """
    
    headers = {"Content-Type": "application/json"}
    data = {
        "prompt": prompt,
        "model": "gemini-1.5-flash-latest",
        "parameters": {
            "temperature": 0.7,
            "maxOutputTokens": 1024,
            "topP": 0.8,
            "topK": 40
        },
        "safetySettings": [
            {"category": "HARM_CATEGORY_DANGEROUS", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
        ]
    }
    
    try:
        response = requests.post(api_url, headers=headers, json=data, timeout=30)
        response.raise_for_status()
        logger.info(f"Gemini API response: {response.json()}")
        return extract_text_from_response(response.json())
    except requests.exceptions.RequestException as e:
        logger.error(f"Gemini API request error: {str(e)}")
        return f"Error generating answer: {str(e)}"
    except Exception as ex:
        logger.error(f"Unexpected error: {str(ex)}")
        return "An unexpected error occurred while generating the answer."

def create_word_document(pdf_path, questions_and_answers):
    """Create enhanced Word document with Q&A"""
    try:
        # Generate output filename
        base_filename = clean_filename(os.path.basename(pdf_path))
        output_filename = f"{base_filename}_QA.docx"
        output_path = os.path.join(PROCESSED_FOLDER, output_filename)
        
        # Create document with enhanced styling
        doc = Document()
        
        # Add title
        title = doc.add_heading('Question & Answer Analysis', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Source PDF: {os.path.basename(pdf_path)}")
        doc.add_paragraph()  # Spacing
        
        # Process each Q&A pair
        for idx, (question, answer) in enumerate(questions_and_answers, 1):
            # Question section
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(f'Question {idx}:')
            question_run.bold = True
            question_run.font.size = Pt(12)
            
            # Question text
            doc.add_paragraph(question)
            
            # Answer section
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run('Answer:')
            answer_run.bold = True
            answer_run.font.size = Pt(12)
            
            # Answer text
            answer_text = doc.add_paragraph(answer or "No answer generated.")
            
            # Add separator
            separator = doc.add_paragraph('_' * 50)
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            doc.add_paragraph()
        
        # Save document
        doc.save(output_path)
        logger.info(f"Successfully created document: {output_filename}")
        return output_path
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        return None

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['file']
        
        # Check if file was selected
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        
        # Validate file type
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Please upload a PDF.'}), 400
        
        try:
            # Save uploaded file
            filename = f"{str(uuid.uuid4())}.pdf"
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            
            # Extract questions
            questions = extract_questions_from_pdf(filepath)
            
            if not questions:
                return jsonify({'error': 'No questions could be extracted from the PDF'}), 400
            
            # Process questions and get answers
            questions_and_answers = []
            for question in questions:
                answer = get_answer_from_gemini(question)
                questions_and_answers.append((question, answer))
            
            # Create Word document
            output_path = create_word_document(file.filename, questions_and_answers)
            
            if output_path and os.path.exists(output_path):
                return send_file(
                    output_path,
                    as_attachment=True,
                    download_name=os.path.basename(output_path)
                )
            else:
                return jsonify({'error': 'Failed to generate document'}), 500
            
        except Exception as e:
            logger.error(f"Processing error: {str(e)}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
        finally:
            # Cleanup uploaded file
            if os.path.exists(filepath):
                os.remove(filepath)
    
    return render_template('index.html')

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File too large. Maximum size is 16MB'}), 413

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8082, debug=True)
