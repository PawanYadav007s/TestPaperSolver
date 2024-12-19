import os
import uuid
import requests
from flask import Flask, render_template, request, send_file
import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# Ensure uploads and processed folders exist
os.makedirs('uploads', exist_ok=True)
os.makedirs('processed', exist_ok=True)

def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return "Could not extract text from PDF"

def extract_text_from_response(api_response):
    try:
        candidates = api_response.get('candidates', [])
        if candidates:
            content = candidates[0].get('content', {})
            parts = content.get('parts', [])
            if parts:
                return parts[0].get('text', "").strip()
        return "No answer generated"
    except (KeyError, IndexError) as e:
        print(f"Error extracting text from response: {e}")
        return "Error in processing API response"

def get_answers_from_gemini(prompt):
    # Replace with your actual API key
    api_key = "AIzaSyDpKtMhIzD-ThtEt6dr4InKQvuQJkcsroc"
    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={api_key}"
    
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]
    }
    
    try:
        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()
        response_data = response.json()
        text_data = extract_text_from_response(response_data)
        return text_data
    except requests.exceptions.RequestException as e:
        print(f"Error calling the Gemini API: {e}")
        return "Error generating answer. Please check the API connection."

def create_word_document(output_path, question_text, answer_text):
    doc = Document()
    
    # Add title
    title = doc.add_heading('Questions and Answers', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    
    # Add question section
    question_heading = doc.add_heading('Question:', level=2)
    question_para = doc.add_paragraph(question_text)
    
    # Add answer section
    answer_heading = doc.add_heading('Answer:', level=2)
    answer_para = doc.add_paragraph(answer_text)
    
    # Add footer with line
    doc.add_paragraph('_' * 50)
    
    doc.save(output_path)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('index.html', error='No file part')
        
        file = request.files['file']
        
        if file.filename == '':
            return render_template('index.html', error='No selected file')
        
        if file and file.filename.lower().endswith('.pdf'):
            # Save uploaded file
            filename = str(uuid.uuid4()) + '.pdf'
            filepath = os.path.join('uploads', filename)
            file.save(filepath)
            
            try:
                # Extract text from PDF
                question_text = extract_text_from_pdf(filepath)
                
                # Get answer from Gemini
                answer_text = get_answers_from_gemini(question_text)
                
                # Create Word document
                output_filename = f"{os.path.splitext(filename)[0]}_QA.docx"
                output_path = os.path.join('processed', output_filename)
                create_word_document(output_path, question_text, answer_text)
                
                return send_file(output_path, as_attachment=True)
            
            except Exception as e:
                return render_template('index.html', error=f'Error processing file: {str(e)}')
        
        return render_template('index.html', error='Invalid file type. Please upload a PDF.')
    
    return render_template('index.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)