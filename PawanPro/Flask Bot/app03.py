import os
import re
import uuid
import requests
from flask import Flask, render_template, request, send_file, jsonify
import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# Ensure uploads and processed folders exist
os.makedirs('uploads', exist_ok=True)
os.makedirs('processed', exist_ok=True)

def extract_questions_from_pdf(pdf_path):
    """Extract individual questions from PDF"""
    questions = []
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = ''
            for page in reader.pages:
                full_text += page.extract_text() + '\n'
        
        # Enhanced regex for question extraction
        question_pattern = r'(\d+[\).][\s]*[^\n]+(?:\n[^\d\n]+)*)'
        raw_questions = re.findall(question_pattern, full_text, re.MULTILINE)
        
        # Clean and filter questions
        for question in raw_questions:
            cleaned_question = re.sub(r'\s+', ' ', question.strip())
            if len(cleaned_question) > 10 and not cleaned_question.lower().startswith(('a.', 'b.', 'c.')):
                questions.append(cleaned_question)
        
        return questions
    except Exception as e:
        app.logger.error(f"Error extracting questions: {e}")
        return []

def extract_text_from_response(api_response):
    """Extract text from Gemini API response"""
    try:
        candidates = api_response.get('candidates', [])
        if candidates:
            content = candidates[0].get('content', {})
            parts = content.get('parts', [])
            if parts:
                return parts[0].get('text', "").strip()
        return "No answer generated."
    except (KeyError, IndexError) as e:
        app.logger.error(f"Error extracting API response: {e}")
        return "Error in processing API response."

def get_answer_from_gemini(question):
    """Get answer from Gemini API"""
    # IMPORTANT: Replace with your actual API key
    API_KEY = "AIzaSyCLq-Y3NHVYGnpwjZx2yIbgl_LiHeYEmO4"
    api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key={API_KEY}"
    
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [
            {"parts": [{"text": f"Provide a comprehensive academic answer to the following question: {question}"}]}
        ],
        "generationConfig": {
            "maxOutputTokens": 1024,
            "temperature": 0.7
        }
    }
    
    try:
        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()
        response_data = response.json()
        return extract_text_from_response(response_data)
    except requests.exceptions.RequestException as e:
        app.logger.error(f"Gemini API error: {e}")
        return f"Error generating answer: {e}"

def create_word_document(pdf_path, questions_and_answers):
    """Create Word document with Q&A"""
    try:
        # Create output folder
        output_folder = 'processed'
        os.makedirs(output_folder, exist_ok=True)
        
        # Generate output filename
        base_filename = os.path.splitext(os.path.basename(pdf_path))[0]
        output_filename = f"{base_filename}_QA.docx"
        output_path = os.path.join(output_folder, output_filename)
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Comprehensive Q&A Document', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process each question-answer pair
        for idx, (question, answer) in enumerate(questions_and_answers, 1):
            # Add space between entries
            if idx > 1:
                doc.add_paragraph()
            
            # Add question
            question_heading = doc.add_heading(f'Question {idx}:', level=2)
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(question)
            question_run.bold = True
            
            # Add answer
            answer_heading = doc.add_heading('Detailed Answer:', level=3)
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run(answer or "No detailed answer generated.")
            
            # Add separator
            doc.add_paragraph('_' * 50)
        
        # Save document
        doc.save(output_path)
        return output_path
    except Exception as e:
        app.logger.error(f"Error creating document: {e}")
        return None

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('index.html', error='No file part')
        
        file = request.files['file']
        
        if file.filename == '':
            return render_template('index.html', error='No selected file')
        
        if file and file.filename.lower().endswith('.pdf'):
            # Save uploaded file
            filename = str(uuid.uuid4()) + '.pdf'
            filepath = os.path.join('uploads', filename)
            file.save(filepath)
            
            try:
                # Extract questions from PDF
                questions = extract_questions_from_pdf(filepath)
                
                if not questions:
                    return render_template('index.html', error='No questions could be extracted from the PDF')
                
                # Process questions and get answers
                questions_and_answers = []
                for question in questions:
                    answer = get_answer_from_gemini(question)
                    questions_and_answers.append((question, answer))
                
                # Create Word document
                output_path = create_word_document(filepath, questions_and_answers)
                
                if output_path:
                    return send_file(output_path, as_attachment=True)
                else:
                    return render_template('index.html', error='Failed to generate document')
            
            except Exception as e:
                app.logger.error(f"Processing error: {e}")
                return render_template('index.html', error=f'Error processing file: {str(e)}')
        
        return render_template('index.html', error='Invalid file type. Please upload a PDF.')
    
    return render_template('index003.html')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)